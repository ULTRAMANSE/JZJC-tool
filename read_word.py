# coding=utf-8
from docx import Document
import re


def word(record_docx):
	"""
	获取记录内容
	:param record_docx:
	:return:
	"""
	doc = Document(record_docx)
	tb = doc.tables
	example_name = []
	description = []
	test_item = []
	results = []
	t_number = 5  # 表格编号
	while t_number < 18:
		temp_a = tb[t_number].cell(0, 2).text
		temp_b = tb[t_number].cell(0, 4).text
		if "用例标识" in temp_a and "用例名称" in temp_b:
			for index, tb_row in enumerate(tb[t_number].column_cells(4)):
				if index is 0:
					continue
				example_name.append(tb_row.text)
				description.append(tb[t_number].column_cells(3)[index].text)
				test_item.append(tb[t_number].column_cells(0)[index].text)
				results.append(tb[t_number].column_cells(5)[index].text)
				c = list(map(list, zip(test_item, description, example_name, results)))
			yield c
			example_name.clear()
			description.clear()
			test_item.clear()
			results.clear()
		t_number += 1


def read_head(docx):
	"""
	返回清单类型
	:param docx:
	:return:
	"""
	doc = Document(docx)
	temp = []
	for p in doc.paragraphs:
		# 获取清单类型
		if p.style.name == 'Heading 2' or p.style.name == '样式2':
			type_test = re.findall(".\d(.+?)用例清单", p.text)
			if type_test:
				temp.append(type_test[0])
	print(temp)
	if temp:
		return temp
	else:
		return False


if __name__ == '__main__':
	read_head("C:\\Users\\ULTRAMA\\Desktop\\JZJC-R-0064-2021 测试记录.docx")
