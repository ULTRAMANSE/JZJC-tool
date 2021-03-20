import sys
from PyQt5.QtWidgets import (QApplication, QWidget, QPushButton, QLabel, QLineEdit, QComboBox, QGridLayout, QFileDialog
, QMainWindow, QStackedWidget)
from PyQt5.Qt import QThread, QMutex
from PyQt5.QtCore import pyqtSignal, pyqtSlot, QFile
from PyQt5.QtGui import QIcon
from docx import Document
import copy
import auto_report
import auto_w
import resource


class WordU(QMainWindow):
	def __init__(self, parent=None):
		super(WordU, self).__init__(parent)
		self.setFixedSize(400, 180)
		self.setWindowTitle("自动填写2.4")
		self.setWindowIcon(QIcon(":/pic/i.ico"))
		self.widget = QWidget()  # 第一页面
		self.second_widget = QWidget()  # 第二页面
		self.third_widget = QWidget()  # 第三页面
		self.file_path = None
		self.save_word = None
		self.file_in = None
		self.save_in = None
		self.style_in = None
		self.yes_b = None
		self.write_word = None  # 多线程

		# 设置stackedWidget
		self.stackedWidget = QStackedWidget()
		self.setCentralWidget(self.stackedWidget)

		# 布局初始化
		self.glayout = QGridLayout()
		self.s_glayout = QGridLayout()
		self.th_glayout = QGridLayout()
		self.glayout.setSpacing(10)

		self.widget.setLayout(self.glayout)
		self.second_widget.setLayout(self.s_glayout)
		self.third_widget.setLayout(self.th_glayout)

		self.stackedWidget.addWidget(self.widget)
		self.stackedWidget.addWidget(self.second_widget)
		self.stackedWidget.addWidget(self.third_widget)

		# 函数初始化
		self.set_prom()
		self.set_menu()
		self.set_second()
		self.set_third()
		self.activity()

	def set_menu(self):
		self.menu_bar = self.menuBar()
		self.menu_bar.setObjectName("menu_bar")
		self.change_recode = self.menu_bar.addAction("说明记录填写")
		self.change = self.menu_bar.addAction("报告填写")
		self.change_flog = self.menu_bar.addAction("标识填写")
		self.menu_bar.addSeparator()

	def activity(self):
		"""
		连接函数
		:return:
		"""
		self.save_in.clicked.connect(self.choose_w_file)
		self.yes_b.clicked.connect(self.start_W)
		self.change_recode.triggered.connect(self.show_fist)
		self.change.triggered.connect(self.show_second)
		self.change_flog.triggered.connect(self.show_third)
		self.read_bu.clicked.connect(self.choose_e_file)
		self.read_in.clicked.connect(self.choose_r_file)
		self.begin_bn.clicked.connect(self.start_in)

		self.word_path.clicked.connect(self.choose_path_file)  # 选择文件
		self.start_sign_in.clicked.connect(self.start_write_sign)  # 开始填写

	def set_prom(self):
		"""
		第一页面布局
		:return:
		"""
		self.save_word = QLineEdit(self)
		self.save_in = QPushButton("选择Word", self)
		self.save_word.setReadOnly(True)
		self.glayout.addWidget(self.save_word, 1, 1, 1, 10)
		self.glayout.addWidget(self.save_in, 1, 11, 1, 4)

		self.style_in = QComboBox(self)
		self.style_in.addItems(["测试说明", "测试记录"])
		self.glayout.addWidget(self.style_in, 2, 4, 1, 2)
		self.yes_b = QPushButton("开始填写", self)
		self.glayout.addWidget(self.yes_b, 2, 8, 1, 4)
		self.prompt = QLabel(self)
		self.glayout.addWidget(self.prompt, 2, 1, 1, 2)

	def set_second(self):
		"""
		第二页面布局
		:return:
		"""
		self.read_excel = QLineEdit(self)
		self.read_bu = QPushButton("选择Excel", self)
		self.s_glayout.addWidget(self.read_excel, 1, 1, 1, 10)
		self.s_glayout.addWidget(self.read_bu, 1, 11, 1, 4)

		self.read_word = QLineEdit(self)
		self.read_in = QPushButton("选择记录", self)
		self.s_glayout.addWidget(self.read_word, 2, 1, 1, 10)
		self.s_glayout.addWidget(self.read_in, 2, 11, 1, 4)

		self.begin_bn = QPushButton("开始生成", self)
		self.s_glayout.addWidget(self.begin_bn, 4, 6, 1, 4)

	def set_third(self):
		self.docx_in = QLineEdit(self)
		self.word_path = QPushButton("选择Word", self)
		self.docx_in.setReadOnly(True)
		self.th_glayout.addWidget(self.docx_in, 1, 1, 1, 10)
		self.th_glayout.addWidget(self.word_path, 1, 11, 1, 4)
		self.start_sign_in = QPushButton("开始填写", self)
		self.th_glayout.addWidget(self.start_sign_in, 2, 6, 1, 4)
		self.display = QLabel(self)
		self.th_glayout.addWidget(self.display, 2, 1, 1, 2)

	def show_fist(self):
		self.stackedWidget.setCurrentIndex(0)

	def show_second(self):
		self.stackedWidget.setCurrentIndex(1)

	def show_third(self):
		self.stackedWidget.setCurrentIndex(2)

	def choose_w_file(self):
		filename, i = QFileDialog.getOpenFileNames(None, "请选择要添加的文件", "./",
												   "Word Files (*.docx);;Word Files (*.doc);;All Files (*)")
		if filename:
			self.save_word.setText(filename[0])

	def choose_e_file(self):
		filename, i = QFileDialog.getOpenFileNames(None, "请选择Excel模板", "./",
												   "Xlsx Files (*.xlsx);;Xls Files (*.xls);;All Files (*)")
		if filename:
			self.read_excel.setText(filename[0])

	def choose_r_file(self):
		filename, i = QFileDialog.getOpenFileNames(None, "请选择记录文件", "./",
												   "Word Files (*.docx);;Word Files (*.doc);;All Files (*)")
		if filename:
			self.read_word.setText(filename[0])

	def choose_path_file(self):
		filename, i = QFileDialog.getOpenFileNames(None, "请选择记录或者说明", "./",
												   "Word Files (*.docx);;Word Files (*.doc);;All Files (*)")
		if filename:
			self.docx_in.setText(filename[0])

	def start_in(self):
		"""
		执行报告填写
		:return:
		"""
		auto_report.auto(self.read_excel.text(), self.read_word.text())

	def start_W(self):
		"""
		自动生成说明记录的表格
		:return:
		"""
		if self.save_word.text() is "":
			self.prompt.setText("请选择文件！！！")
		else:
			self.prompt.setText("执行中...")
			in_item = self.style_in.currentIndex()
			self.write_word = ThreadRW(self.save_word.text(), in_item)
			self.write_word.str_out.connect(self.prompt_out)
			self.write_word.start()

	def start_write_sign(self):
		auto_w.read(self.docx_in.text())

	@pyqtSlot(str)
	def prompt_out(self, i):
		self.prompt.setText(i)


lock = QMutex()


class ThreadRW(QThread):
	str_out = pyqtSignal(str)

	def __init__(self, wfile=None, item=None, parent=None):
		super().__init__(parent)
		self.W_file = wfile  # word文件名
		self.Item = item  # 类型选项
		self.number = 0  # 记录差
		self.num = 0  # 列表下表值
		self.t_number = 5  # 表格index
		self.t_number_copy = 0  # 复制表格index
		self.t_rows = 1  # 表格的行index
		self.example_name = []  # 用例名称
		self.identity = []  # 用例标识

	def move_table_after(self, table, paragraph):
		tbl, p = table._tbl, paragraph._p
		p.addnext(tbl)

	def run(self):
		lock.lock()
		if self.Item == 1:
			self.number = 2
		doc = Document(self.W_file)
		tb = doc.tables

		while True:
			temp_a = tb[self.t_number].cell(0, 2).text  # 临时变量，用于查找文档符合条件的用例清单表
			temp_b = tb[self.t_number].cell(0, 4).text
			if "用例标识" in temp_a and "用例名称" in temp_b:
				self.t_number_copy = self.t_number
				break
			self.t_number += 1

		for index, tb_row in enumerate(tb[self.t_number_copy].column_cells(2)):
			if index is 0:
				continue
			self.example_name.append(tb[self.t_number_copy].column_cells(4)[index].text)
			self.identity.append(tb_row.text)

		while True:
			a = tb[self.t_number].cell(0, 0).text
			b = tb[self.t_number].cell(0, 4 + self.Item).text
			if "用例名称" in a and "用例标识" in b:
				self.t_number_copy = self.t_number
				break
			self.t_number += 1

		while len(self.example_name) - 1 >= self.num:
			if self.t_number_copy == 0:
				break
			copy_tb = copy.deepcopy(tb[self.t_number_copy])
			tb[self.t_number_copy].cell(0, 1).text = self.example_name[self.num]
			run = tb[self.t_number_copy].cell(0, 5 + self.number).paragraphs[0].add_run(self.identity[self.num])
			run.font.name = 'Times New Roman'
			self.t_number_copy += 1
			try:
				self.num += 1
				if len(self.example_name) > self.num:
					pg = doc.paragraphs
					self.move_table_after(copy_tb, pg[len(pg) - 1])
					doc.add_paragraph(" ")
					doc.add_paragraph(" ")
					tb = doc.tables
			except BaseException as e:
				print(e, "异常")
				break

		doc.save(self.W_file)
		re_str = "执行完成"
		self.str_out.emit(str(re_str))
		lock.unlock()


# class CommonHelper:
# 	def __init__(self):
# 		pass
#
# 	@staticmethod
# 	def read_qss(stylefile):
# 		with open(stylefile, 'r') as f:
# 			return f.read()


if __name__ == '__main__':
	App = QApplication(sys.argv)
	qss_file = QFile(":/qss/auto-type.qss")
	qss_file.open(QFile.ReadOnly)
	qss = str(qss_file.readAll(), encoding='utf-8')
	qss_file.close()
	# style = CommonHelper.read_qss("./auto-type.qss")
	ex = WordU()
	ex.setStyleSheet(qss)
	ex.show()
	sys.exit(App.exec_())
